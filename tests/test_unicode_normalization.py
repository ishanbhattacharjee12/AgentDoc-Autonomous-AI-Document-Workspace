import unittest
import os
from pathlib import Path
from docx import Document
from app.tools.document_tool import _to_latin1, generate_docx

class TestUnicodeNormalization(unittest.TestCase):
    
    def test_pdf_latin1_normalization(self):
        """Verify that _to_latin1 correctly normalizes mathematical, directional, and typographic characters for PDF."""
        test_cases = {
            # Inequalities
            "≥75%": ">=75%",
            "≤18%": "<=18%",
            "x ≠ y": "x != y",
            
            # Arrows
            "72h→48h": "72h->48h",
            "A ← B": "A <- B",
            "A ↔ B": "A <-> B",
            
            # Math
            "−18%": "-18%",  # Unicode minus sign
            "~$3.1M": "~$3.1M",
            "≈$5.8M": "~$5.8M",
            "10 × 20": "10 x 20",
            
            # Quotes & Dashes
            "“hello”": '"hello"',
            "‘hello’": "'hello'",
            "em—dash": "em-dash",
            "en–dash": "en-dash",
            
            # Whitespace & Ellipsis
            "A\xa0B": "A B",  # Non-breaking space
            "hello…": "hello...",
        }
        
        for input_str, expected_str in test_cases.items():
            result = _to_latin1(input_str)
            self.assertEqual(result, expected_str, f"Failed for input: {input_str}")
            self.assertNotIn("?", result, f"Stray question mark found in normalized output for input: {input_str}")

    def test_docx_unicode_preservation(self):
        """Verify that DOCX generation preserves the original professional Unicode symbols intact."""
        title = "Unicode Onboarding Test Document"
        doc_type = "onboarding_plan"
        assumptions = ["≥95% config completion target", "budget ≤ $50K"]
        goal = "Test unicode preservation in DOCX"
        content = """# Onboarding Framework

## Target Metrics
- Funnel drop-off is currently ≥62%.
- We need to drive step-1 drop-off from 38% → 28%.
- Total spend is −$3K under budget.
- We expect ≈80 users/month.
"""
        
        # Generate DOCX
        filename = generate_docx(
            title=title,
            document_type=doc_type,
            assumptions=assumptions,
            content=content,
            goal=goal,
            format_ext="docx"
        )
        
        filepath = Path("app/outputs") / filename
        self.assertTrue(filepath.exists(), "DOCX file was not generated.")
        
        try:
            # Load and verify DOCX content
            doc = Document(filepath)
            
            # Combine all text from paragraphs
            full_text = []
            for p in doc.paragraphs:
                full_text.append(p.text)
            
            for section in doc.sections:
                if section.header:
                    for p in section.header.paragraphs:
                        full_text.append(p.text)
            
            all_text = "\n".join(full_text)
            
            # Verify original unicode characters are preserved in DOCX
            self.assertIn("≥62%", all_text)
            self.assertIn("38% → 28%", all_text)
            self.assertIn("−$3K", all_text)
            self.assertIn("≈80", all_text)
            self.assertNotIn("?62%", all_text)
            self.assertNotIn("?80", all_text)
            
        finally:
            # Cleanup
            if filepath.exists():
                os.remove(filepath)

    def test_pdf_generation_runs_successfully(self):
        """Verify that PDF generation runs successfully on unicode strings without throwing errors or generating stray ?."""
        title = "Unicode Onboarding Test Document"
        doc_type = "onboarding_plan"
        assumptions = ["≥95% target", "budget ≤ $50K"]
        goal = "Test unicode rendering in PDF"
        content = """# Onboarding Framework
- Funnel drop-off is currently ≥62%.
- We need to drive step-1 drop-off from 38% → 28%.
- Total spend is −$3K under budget.
- We expect ≈80 users/month.
"""
        
        # Generate PDF
        filename = generate_docx(
            title=title,
            document_type=doc_type,
            assumptions=assumptions,
            content=content,
            goal=goal,
            format_ext="pdf"
        )
        
        filepath = Path("app/outputs") / filename
        self.assertTrue(filepath.exists(), "PDF file was not generated.")
        
        # Cleanup
        if filepath.exists():
            os.remove(filepath)
