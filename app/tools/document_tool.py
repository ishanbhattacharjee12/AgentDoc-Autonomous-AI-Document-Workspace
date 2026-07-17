"""Document tool for deterministic DOCX generation.

Uses python-docx to create professional Word documents
from synthesized content.
"""

import logging
import re
from datetime import datetime
from pathlib import Path
import markdown
from fpdf import FPDF
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import time
from app.config import OUTPUT_DIR
from app.llm.budget_guard import check_budget_pdf

logger = logging.getLogger(__name__)


def sanitize_filename(name: str) -> str:
    """Create a safe filename from a string."""
    # Remove unsafe characters
    safe = re.sub(r'[^\w\s-]', '', name.lower())
    safe = re.sub(r'[\s]+', '_', safe.strip())
    # Truncate to reasonable length
    return safe[:50] if safe else "document"


class DOCXTheme:
    # Centralized colors (HEX and RGB format)
    HEX_PRIMARY = "1B3F32"       # Dark Forest Green
    HEX_SECONDARY = "5C6B62"     # Slate / Medium Sage Grey
    HEX_TEXT = "282C2A"          # Off-black charcoal
    HEX_MUTED = "888C8A"         # Muted grey-sage
    HEX_DIVIDER = "DAE0DC"       # Light grey-green
    HEX_BG_LIGHT = "F4F6F4"      # Clean background cover fill / standard shading
    HEX_BG_SAGE = "EDF3EF"       # Executive summary soft sage box fill

    RGB_PRIMARY = RGBColor(27, 63, 50)
    RGB_SECONDARY = RGBColor(92, 107, 98)
    RGB_TEXT = RGBColor(40, 44, 42)
    RGB_MUTED = RGBColor(136, 140, 138)

    # Fonts
    FONT_BODY = "Arial"
    FONT_TITLE = "Times New Roman"

    # Sizes (pt)
    SIZE_TITLE = 26
    SIZE_SUBTITLE = 11
    SIZE_H1 = 15
    SIZE_H2 = 12.5
    SIZE_H3 = 11
    SIZE_BODY = 10
    SIZE_FOOTER = 8.5


# --- OXML HELPER FUNCTIONS FOR PYTHON-DOCX ---

def _set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is not None:
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is not None:
        tcPr.remove(tcMar)
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_cell_borders(cell, **kwargs):
    """
    Sets borders for a cell.
    kwargs: top, bottom, left, right
    value: dict like {'sz': 12, 'val': 'single', 'color': 'HEX', 'space': '0'}
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            el = tcBorders.find(qn(tag))
            if el is not None:
                tcBorders.remove(el)
            element = OxmlElement(tag)
            for key, val in edge_data.items():
                element.set(qn(f'w:{key}'), str(val))
            tcBorders.append(element)


def _set_paragraph_bottom_border(p, hex_color="1B3F32", size=12):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = pBdr.find(qn('w:bottom'))
    if bottom is not None:
        pBdr.remove(bottom)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '8')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)


def _set_paragraph_top_border(p, hex_color="DAE0DC", size=4):
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    top = pBdr.find(qn('w:top'))
    if top is not None:
        pBdr.remove(top)
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size))
    top.set(qn('w:space'), '12')
    top.set(qn('w:color'), hex_color)
    pBdr.append(top)


def _prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = trPr.find(qn('w:cantSplit'))
    if cantSplit is None:
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)


def _set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = trPr.find(qn('w:tblHeader'))
    if tblHeader is None:
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)


def _add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


# --- REUSABLE DOCX COMPONENT RENDERING HELPERS ---

def _docx_build_cover_page(doc: Document, title: str, document_type: str, gen_date: str) -> None:
    cover_section = doc.sections[0]
    cover_section.top_margin = Inches(1.0)
    cover_section.bottom_margin = Inches(1.0)
    cover_section.left_margin = Inches(1.0)
    cover_section.right_margin = Inches(1.0)

    # 1x1 table acting as cover page card
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    _set_cell_background(cell, DOCXTheme.HEX_BG_LIGHT)
    _set_cell_margins(cell, top=1440, bottom=1440, left=720, right=720)
    _set_cell_borders(
        cell,
        left={'val': 'single', 'sz': '108', 'color': DOCXTheme.HEX_PRIMARY, 'space': '0'},
        top={'val': 'none'},
        right={'val': 'none'},
        bottom={'val': 'none'}
    )
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    
    # Nested table for monogram brand logo and wordmark
    logo_table = cell.add_table(rows=1, cols=2)
    logo_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    cell_logo = logo_table.cell(0, 0)
    cell_logo.width = Inches(0.5)
    _set_cell_background(cell_logo, DOCXTheme.HEX_PRIMARY)
    _set_cell_margins(cell_logo, top=60, bottom=60, left=60, right=60)
    _set_cell_borders(cell_logo, top={'val':'none'}, left={'val':'none'}, right={'val':'none'}, bottom={'val':'none'})
    logo_p = cell_logo.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_p.add_run("AD")
    logo_run.bold = True
    logo_run.font.name = DOCXTheme.FONT_BODY
    logo_run.font.size = Pt(9.5)
    logo_run.font.color.rgb = RGBColor(255, 255, 255)
    
    cell_brand = logo_table.cell(0, 1)
    cell_brand.width = Inches(5.0)
    _set_cell_margins(cell_brand, top=60, bottom=60, left=120, right=60)
    _set_cell_borders(cell_brand, top={'val':'none'}, left={'val':'none'}, right={'val':'none'}, bottom={'val':'none'})
    brand_p = cell_brand.paragraphs[0]
    brand_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    brand_run = brand_p.add_run("AgentDoc AI Workspace")
    brand_run.bold = True
    brand_run.font.name = DOCXTheme.FONT_BODY
    brand_run.font.size = Pt(10.5)
    brand_run.font.color.rgb = DOCXTheme.RGB_PRIMARY
    
    # Large spacers
    p_space = cell.add_paragraph()
    p_space.paragraph_format.space_before = Pt(72)
    p_space.paragraph_format.space_after = Pt(0)
    
    # Document Type
    p_type = cell.add_paragraph()
    p_type.paragraph_format.space_before = Pt(0)
    p_type.paragraph_format.space_after = Pt(6)
    run_type = p_type.add_run(document_type.upper().replace('_', ' '))
    run_type.bold = True
    run_type.font.name = DOCXTheme.FONT_BODY
    run_type.font.size = Pt(10.5)
    run_type.font.color.rgb = DOCXTheme.RGB_SECONDARY
    
    # Title (Serif bold Times New Roman size 26)
    p_title = cell.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(16)
    p_title.paragraph_format.line_spacing = 1.1
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.font.name = DOCXTheme.FONT_TITLE
    run_title.font.size = Pt(26)
    run_title.font.color.rgb = DOCXTheme.RGB_PRIMARY
    _set_paragraph_bottom_border(p_title, hex_color=DOCXTheme.HEX_PRIMARY, size=12)
    
    # Subtitle
    p_sub = cell.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(12)
    p_sub.paragraph_format.space_after = Pt(96)
    run_sub = p_sub.add_run("Prepared for: Strategic Executive Stakeholders")
    run_sub.italic = True
    run_sub.font.name = DOCXTheme.FONT_BODY
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = DOCXTheme.RGB_TEXT
    
    # Metadata Row at bottom (3 columns)
    meta_table = cell.add_table(rows=2, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_w = Inches(1.8)
    for col in meta_table.columns:
        col.width = col_w
        
    labels = ["CLASSIFICATION", "PREPARED BY", "DATE OF ISSUE"]
    vals = ["Confidential / Executive", "AgentDoc Services", gen_date]
    
    for i, label in enumerate(labels):
        cell_lbl = meta_table.cell(0, i)
        cell_lbl.width = col_w
        _set_cell_margins(cell_lbl, top=0, bottom=0, left=0, right=0)
        _set_cell_borders(cell_lbl, top={'val':'none'}, left={'val':'none'}, right={'val':'none'}, bottom={'val':'none'})
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(2)
        run = p_lbl.add_run(label)
        run.bold = True
        run.font.name = DOCXTheme.FONT_BODY
        run.font.size = Pt(7.5)
        run.font.color.rgb = DOCXTheme.RGB_SECONDARY
        
    for i, val in enumerate(vals):
        cell_val = meta_table.cell(1, i)
        cell_val.width = col_w
        _set_cell_margins(cell_val, top=0, bottom=0, left=0, right=0)
        _set_cell_borders(cell_val, top={'val':'none'}, left={'val':'none'}, right={'val':'none'}, bottom={'val':'none'})
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(0)
        run = p_val.add_run(val)
        run.font.name = DOCXTheme.FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = DOCXTheme.RGB_PRIMARY


def _docx_build_running_header(section, title: str, document_type: str) -> None:
    header = section.header
    tbl = header.add_table(rows=1, cols=2, width=Inches(6.5))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    c1 = tbl.cell(0, 0)
    c1.width = Inches(3.25)
    _set_cell_margins(c1, top=60, bottom=60, left=0, right=60)
    _set_cell_borders(c1, left={'val':'none'}, top={'val':'none'}, right={'val':'none'}, bottom={'val':'single', 'sz':'4', 'color': DOCXTheme.HEX_DIVIDER})
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(title)
    r1.font.name = DOCXTheme.FONT_BODY
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = DOCXTheme.RGB_MUTED
    
    c2 = tbl.cell(0, 1)
    c2.width = Inches(3.25)
    _set_cell_margins(c2, top=60, bottom=60, left=60, right=0)
    _set_cell_borders(c2, left={'val':'none'}, top={'val':'none'}, right={'val':'none'}, bottom={'val':'single', 'sz':'4', 'color': DOCXTheme.HEX_DIVIDER})
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(document_type.upper().replace('_', ' '))
    r2.font.name = DOCXTheme.FONT_BODY
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = DOCXTheme.RGB_MUTED


def _docx_build_running_footer(section) -> None:
    footer = section.footer
    tbl = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    c1 = tbl.cell(0, 0)
    c1.width = Inches(4.5)
    _set_cell_margins(c1, top=60, bottom=0, left=0, right=0)
    _set_cell_borders(c1, left={'val':'none'}, top={'val':'none'}, right={'val':'none'}, bottom={'val':'none'})
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run("AgentDoc AI Workspace — Autonomous Document Generation Agent")
    r1.italic = True
    r1.font.name = DOCXTheme.FONT_BODY
    r1.font.size = Pt(8)
    r1.font.color.rgb = DOCXTheme.RGB_MUTED
    
    c2 = tbl.cell(0, 1)
    c2.width = Inches(2.0)
    _set_cell_margins(c2, top=60, bottom=0, left=0, right=0)
    _set_cell_borders(c2, left={'val':'none'}, top={'val':'none'}, right={'val':'none'}, bottom={'val':'none'})
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("Page ")
    r2.font.name = DOCXTheme.FONT_BODY
    r2.font.size = Pt(8)
    r2.font.color.rgb = DOCXTheme.RGB_MUTED
    _add_page_number(p2.add_run())


def _docx_render_callout_box(doc: Document, text: str, hex_border_color: str, hex_fill_color: str, text_style: str = "") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    _set_cell_background(cell, hex_fill_color)
    _set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    _set_cell_borders(
        cell,
        left={'val': 'single', 'sz': '36', 'color': hex_border_color, 'space': '0'},
        top={'val': 'none'},
        right={'val': 'none'},
        bottom={'val': 'none'}
    )
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
            
        run.font.name = DOCXTheme.FONT_BODY
        run.font.size = Pt(DOCXTheme.SIZE_BODY)
        run.font.color.rgb = DOCXTheme.RGB_TEXT
        if "I" in text_style:
            run.italic = True
            
    _prevent_row_split(table.rows[0])
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(6)


def _docx_render_executive_summary_panel(doc: Document, blocks: list) -> None:
    lines = []
    for b in blocks:
        if b["type"] == "paragraph":
            lines.extend([l.strip() for l in b["content"].split("\n") if l.strip()])
        elif b["type"] == "list":
            items = parse_list_items(b["content"])
            for item in items:
                prefix = f"{item['marker']}. " if item["is_number"] else "• "
                lines.append(f"{prefix}{item['text']}")
                
    if not lines:
        return
        
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    
    _set_cell_background(cell, DOCXTheme.HEX_BG_SAGE)
    _set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    _set_cell_borders(
        cell,
        left={'val': 'single', 'sz': '36', 'color': DOCXTheme.HEX_PRIMARY, 'space': '0'},
        top={'val': 'none'},
        right={'val': 'none'},
        bottom={'val': 'none'}
    )
    
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            p.text = ""
            first = False
        else:
            p = cell.add_paragraph()
            
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = DOCXTheme.FONT_BODY
            run.font.size = Pt(DOCXTheme.SIZE_BODY)
            run.font.color.rgb = DOCXTheme.RGB_TEXT
            run.italic = True
            
    _prevent_row_split(table.rows[0])
    
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(12)


def _docx_render_consulting_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
        
    num_cols = max(len(r) for r in rows)
    if num_cols == 0:
        return
        
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    col_w = Inches(6.5 / num_cols)
    for col in table.columns:
        col.width = col_w
        
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        _prevent_row_split(row)
        if r_idx == 0:
            _set_repeat_header(row)
            
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = row.cells[c_idx]
                cell.width = col_w
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                
                _set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                
                if r_idx == 0:
                    _set_cell_background(cell, DOCXTheme.HEX_PRIMARY)
                    _set_cell_borders(
                        cell,
                        top={'val':'none'}, left={'val':'none'}, right={'val':'none'},
                        bottom={'val': 'single', 'sz': '8', 'color': DOCXTheme.HEX_PRIMARY}
                    )
                else:
                    if r_idx % 2 == 0:
                        _set_cell_background(cell, DOCXTheme.HEX_BG_LIGHT)
                    _set_cell_borders(
                        cell,
                        top={'val':'none'}, left={'val':'none'}, right={'val':'none'},
                        bottom={'val': 'single', 'sz': '4', 'color': DOCXTheme.HEX_DIVIDER}
                    )
                    
                text = _clean_markdown(cell_text)
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                        
                    run.font.name = DOCXTheme.FONT_BODY
                    run.font.size = Pt(9)
                    if r_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = DOCXTheme.RGB_TEXT
                        
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(8)


def _docx_render_risk_matrix(doc: Document, block_content: str) -> None:
    items = parse_list_items(block_content)
    table_rows = [["Risk Event / Vulnerability", "Mitigation Strategy"]]
    
    for item in items:
        text = item["text"]
        parts = re.split(r'[:\-]', text, maxsplit=1)
        if len(parts) >= 2:
            risk = parts[0].strip().replace('**', '').replace('*', '')
            mitigation = parts[1].strip()
            table_rows.append([f"**Risk:** {risk}", mitigation])
        else:
            table_rows.append(["Key Vulnerability Identified", text])
            
    _docx_render_consulting_table(doc, table_rows)


def _docx_render_phase_card(doc: Document, phase_title: str, lines: list[str]) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_header = table.cell(0, 0)
    cell_body = table.cell(1, 0)
    
    cell_header.width = Inches(6.5)
    cell_body.width = Inches(6.5)
    
    _prevent_row_split(table.rows[0])
    _prevent_row_split(table.rows[1])
    
    # Header styling
    _set_cell_background(cell_header, DOCXTheme.HEX_PRIMARY)
    _set_cell_margins(cell_header, top=80, bottom=80, left=150, right=150)
    _set_cell_borders(cell_header, top={'val':'single', 'sz':'4', 'color': DOCXTheme.HEX_DIVIDER}, left={'val':'single', 'sz':'4', 'color': DOCXTheme.HEX_DIVIDER}, right={'val':'single', 'sz':'4', 'color': DOCXTheme.HEX_DIVIDER}, bottom={'val':'none'})
    
    p_head = cell_header.paragraphs[0]
    p_head.paragraph_format.space_before = Pt(0)
    p_head.paragraph_format.space_after = Pt(0)
    run_head = p_head.add_run(phase_title.upper())
    run_head.bold = True
    run_head.font.name = DOCXTheme.FONT_BODY
    run_head.font.size = Pt(9.5)
    run_head.font.color.rgb = RGBColor(255, 255, 255)
    
    # Body styling
    _set_cell_background(cell_body, DOCXTheme.HEX_BG_LIGHT)
    _set_cell_margins(cell_body, top=120, bottom=120, left=150, right=150)
    _set_cell_borders(cell_body, top={'val':'none'}, left={'val':'single', 'sz':'4', 'color': DOCXTheme.HEX_DIVIDER}, right={'val':'single', 'sz':'4', 'color': DOCXTheme.HEX_DIVIDER}, bottom={'val':'single', 'sz':'4', 'color': DOCXTheme.HEX_DIVIDER})
    
    first_item = True
    for line in lines:
        stripped = line.strip().replace('**', '').replace('*', '')
        parts = re.split(r'^(\*?\*?Objective:\*?\*?|\*?\*?Deliverables:\*?\*?|\*?\*?Timeline:\*?\*?|\*?\*?Budget:\*?\*?)', stripped, flags=re.IGNORECASE)
        
        if first_item:
            p_item = cell_body.paragraphs[0]
            p_item.text = ""
            first_item = False
        else:
            p_item = cell_body.add_paragraph()
            
        p_item.paragraph_format.space_before = Pt(2)
        p_item.paragraph_format.space_after = Pt(2)
        p_item.paragraph_format.line_spacing = 1.15
        
        if len(parts) >= 3:
            label_part = parts[1].strip(" *:")
            body_part = parts[2].strip()
            
            run_lbl = p_item.add_run(f"•  {label_part}: ")
            run_lbl.bold = True
            run_lbl.font.name = DOCXTheme.FONT_BODY
            run_lbl.font.size = Pt(9.5)
            run_lbl.font.color.rgb = DOCXTheme.RGB_PRIMARY
            
            run_val = p_item.add_run(body_part)
            run_val.font.name = DOCXTheme.FONT_BODY
            run_val.font.size = Pt(9.5)
            run_val.font.color.rgb = DOCXTheme.RGB_TEXT
        else:
            run_val = p_item.add_run(f"•  {stripped}")
            run_val.font.name = DOCXTheme.FONT_BODY
            run_val.font.size = Pt(9.5)
            run_val.font.color.rgb = DOCXTheme.RGB_TEXT
            
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(8)


# --- MAIN ORCHESTRATION PIPELINE ---

def generate_docx(
    title: str,
    document_type: str,
    assumptions: list[str],
    content: str,
    goal: str,
    format_ext: str = "docx"
) -> str:
    """Generate a professional document in the requested format.

    Args:
        title: Document title
        document_type: Type of document being generated
        assumptions: List of assumptions made
        content: The synthesized document content
        goal: The interpreted goal

    Returns:
        The filename of the generated document.
    """
    content = _editorial_cleanup_pass(content)
    OUTPUT_DIR.mkdir(exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_type = sanitize_filename(document_type)
    
    fmt = format_ext.lower().strip('.')
    if fmt not in ["docx", "pdf", "html", "md", "markdown"]:
        fmt = "docx"
        
    filename = f"agentdoc_{safe_type}_{timestamp}.{fmt}"
    filepath = OUTPUT_DIR / filename

    start_build_time = time.time()
    
    if fmt in ["md", "markdown"]:
        res = _generate_markdown(filepath, title, document_type, assumptions, content, goal, filename)
        check_budget_pdf(time.time() - start_build_time, fmt)
        return res
    elif fmt == "html":
        res = _generate_html(filepath, title, document_type, assumptions, content, goal, filename)
        check_budget_pdf(time.time() - start_build_time, fmt)
        return res
    elif fmt == "pdf":
        res = _generate_pdf(filepath, title, document_type, assumptions, content, goal, filename)
        check_budget_pdf(time.time() - start_build_time, fmt)
        return res

    doc = Document()
    gen_date = datetime.now().strftime("%B %d, %Y")

    # Extract H1 Title from content if present
    extracted_title = ""
    for line in content.split('\n'):
        line_strip = line.strip()
        h1_match = re.match(r'^#\s+(.+)', line_strip)
        h1_alt_match = re.match(r'^##\s+(.+)', line_strip)
        if h1_match or h1_alt_match:
            match_obj = h1_match if h1_match else h1_alt_match
            extracted_title = match_obj.group(1).strip().replace('**', '').replace('*', '')
            break
            
    if extracted_title:
        title = extracted_title
    else:
        clean_type = document_type.replace('_', ' ').title()
        if len(title) > 60 or "." in title or "," in title:
            lower_title = title.lower()
            if "onboarding" in lower_title:
                title = f"Customer Onboarding {clean_type}"
            elif "activation" in lower_title:
                title = f"Customer Activation {clean_type}"
            elif "experience" in lower_title:
                title = f"Customer Experience {clean_type}"
            elif "retention" in lower_title:
                title = f"Customer Retention {clean_type}"
            else:
                words = [w for w in title.split() if w.lower() not in ["a", "an", "the", "we", "need", "to", "create", "generate", "write", "develop", "provide", "design", "draft"]]
                short_title = " ".join(words[:4]).strip(" ,.-/\\")
                title = f"{short_title.title()} {clean_type}"

    # Set normal styles
    style = doc.styles['Normal']
    font = style.font
    font.name = DOCXTheme.FONT_BODY
    font.size = Pt(DOCXTheme.SIZE_BODY)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    # 1. BUILD COVER PAGE (Section 1)
    _docx_build_cover_page(doc, title, document_type, gen_date)

    # 2. MAIN CONTENT (Section 2)
    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body_section.top_margin = Inches(1.0)
    body_section.bottom_margin = Inches(1.0)
    body_section.left_margin = Inches(1.0)
    body_section.right_margin = Inches(1.0)
    
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False
    
    _docx_build_running_header(body_section, title, document_type)
    _docx_build_running_footer(body_section)

    # 3. RENDER CONTENT SECTIONS
    _add_content_sections(doc, content)

    doc.save(str(filepath))
    logger.info("Generated DOCX: %s", filename)
    check_budget_pdf(time.time() - start_build_time, fmt)
    return filename


def _generate_markdown(filepath: Path, title: str, document_type: str, assumptions: list[str], content: str, goal: str, filename: str) -> str:
    lines = [
        f"# {title}",
        f"**Document Type:** {document_type.replace('_', ' ').title()}",
        f"**Generated:** {datetime.now().strftime('%B %d, %Y')}",
        ""
    ]
    if goal:
        lines.extend(["## Executive Summary", goal, ""])
    if assumptions:
        lines.extend(["## Key Assumptions"])
        lines.extend([f"- {a}" for a in assumptions])
        lines.extend([""])
        
    lines.append(content)
    lines.extend([
        "",
        "---",
        "*Generated by AgentDoc — Autonomous Document Generation Agent*"
    ])
    
    filepath.write_text("\n".join(lines), encoding="utf-8")
    return filename


def _generate_html(filepath: Path, title: str, document_type: str, assumptions: list[str], content: str, goal: str, filename: str) -> str:
    md_text = _generate_markdown(filepath, title, document_type, assumptions, content, goal, filename)
    # Read the markdown back
    md_content = filepath.read_text(encoding="utf-8")
    
    html_body = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
    body {{ font-family: 'Segoe UI', system-ui, sans-serif; line-height: 1.6; color: #333; max-width: 900px; margin: 40px auto; padding: 0 20px; }}
    h1 {{ color: #2c3e50; border-bottom: 2px solid #eee; padding-bottom: 10px; }}
    h2 {{ color: #34495e; margin-top: 30px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
    th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
    th {{ background-color: #f8f9fa; font-weight: 600; }}
    hr {{ border: 0; border-top: 1px solid #eee; margin: 40px 0; }}
    .footer {{ color: #888; font-size: 0.9em; text-align: center; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""

    filepath.write_text(html, encoding="utf-8")
    return filename


def _to_latin1(text: str) -> str:
    """Sanitize unicode string to latin-1, replacing unsupported characters with equivalents."""
    replacements = {
        '\u2014': '-',   # em-dash
        '\u2013': '-',   # en-dash
        '\u201c': '"',   # left double quote
        '\u201d': '"',   # right double quote
        '\u2018': "'",   # left single quote
        '\u2019': "'",   # right single quote
        '\u2022': '\x95', # bullet point
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)
    return text.encode('latin-1', 'replace').decode('latin-1')


class PDFTheme:
    # Margins (mm)
    MARGIN_LEFT = 25.4
    MARGIN_RIGHT = 25.4
    MARGIN_TOP = 25.4
    MARGIN_BOTTOM = 25.4
    
    # Grayscale-friendly Sage & Green Palette (RGB)
    COLOR_PRIMARY = (27, 63, 50)       # Dark Forest Green spine & main headers
    COLOR_SECONDARY = (92, 107, 98)    # Slate / Medium Sage Grey for secondary headers & subtitles
    COLOR_TEXT = (40, 44, 42)          # Off-black charcoal for body text
    COLOR_MUTED = (136, 140, 138)      # Muted grey-sage for headers/footers
    COLOR_DIVIDER = (218, 224, 220)    # Light grey-green for dividers & borders
    
    # Table Styling
    COLOR_TABLE_HEADER_FILL = (27, 63, 50)    # Forest Green Table Header
    COLOR_TABLE_HEADER_TEXT = (255, 255, 255)   # White Header Text
    COLOR_TABLE_CELL_FILL_EVEN = (243, 245, 244) # Very soft gray-green background rows
    COLOR_TABLE_BORDER = (218, 224, 220)       # Light border lines
    
    # Typography
    FONT_FAMILY = "helvetica"
    
    # Font Sizes (pt)
    SIZE_TITLE = 26
    SIZE_SUBTITLE = 12
    SIZE_H1 = 15
    SIZE_H2 = 12.5
    SIZE_H3 = 11
    SIZE_BODY = 10
    SIZE_FOOTER = 8
    
    # Line Heights (mm)
    LEAD_TITLE = 11
    LEAD_SUBTITLE = 8
    LEAD_H1 = 8
    LEAD_H2 = 7
    LEAD_H3 = 6
    LEAD_BODY = 6
    LEAD_FOOTER = 5.5


class PDF(FPDF):
    def __init__(self, doc_title="", doc_type="", gen_date="", *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.doc_title = doc_title
        self.doc_type = doc_type
        self.gen_date = gen_date
        # Enable alias for total page count placeholder {nb}
        self.alias_nb_pages()

    def check_page_break(self, height):
        """Force a page break if remaining vertical space is less than height."""
        if self.get_y() + height > self.h - self.b_margin:
            self.add_page()

    def header(self):
        # Header runs only on content pages (page > 1)
        if self.page_no() > 1:
            current_l_margin = self.l_margin
            current_r_margin = self.r_margin
            self.set_left_margin(PDFTheme.MARGIN_LEFT)
            self.set_right_margin(PDFTheme.MARGIN_RIGHT)
            
            self.set_y(12)
            self.set_font(PDFTheme.FONT_FAMILY, "I", PDFTheme.SIZE_FOOTER)
            self.set_text_color(*PDFTheme.COLOR_MUTED)
            
            # Left Header: title & type
            clean_type = self.doc_type.replace('_', ' ').title()
            header_text = f"AgentDoc | {self.doc_title} ({clean_type})"
            self.cell(0, 8, _to_latin1(header_text), align="L")
            
            # Right Header: Date
            self.set_x(self.w - PDFTheme.MARGIN_RIGHT - 50)
            self.cell(50, 8, _to_latin1(self.gen_date), align="R")
            
            # Draw header line divider
            line_y = 20
            self.set_draw_color(*PDFTheme.COLOR_DIVIDER)
            self.set_line_width(0.3)
            self.line(PDFTheme.MARGIN_LEFT, line_y, self.w - PDFTheme.MARGIN_RIGHT, line_y)
            
            self.set_left_margin(current_l_margin)
            self.set_right_margin(current_r_margin)
            self.set_y(PDFTheme.MARGIN_TOP)

    def footer(self):
        # Footer runs only on content pages (page > 1)
        if self.page_no() > 1:
            current_l_margin = self.l_margin
            current_r_margin = self.r_margin
            self.set_left_margin(PDFTheme.MARGIN_LEFT)
            self.set_right_margin(PDFTheme.MARGIN_RIGHT)
            
            # Draw footer line divider
            line_y = self.h - 18
            self.set_draw_color(*PDFTheme.COLOR_DIVIDER)
            self.set_line_width(0.3)
            self.line(PDFTheme.MARGIN_LEFT, line_y, self.w - PDFTheme.MARGIN_RIGHT, line_y)
            
            self.set_y(-15)
            self.set_font(PDFTheme.FONT_FAMILY, "I", PDFTheme.SIZE_FOOTER)
            self.set_text_color(*PDFTheme.COLOR_MUTED)
            
            self.cell(0, 10, _to_latin1("Generated by AgentDoc Services"), align="L")
            self.set_x(self.w - PDFTheme.MARGIN_RIGHT - 40)
            self.cell(40, 10, _to_latin1(f"Page {self.page_no()} of {{nb}}"), align="R")
            
            self.set_left_margin(current_l_margin)
            self.set_right_margin(current_r_margin)


# ---------------- VECTOR GRAPHIC ICONS next to HEADINGS ----------------
def draw_heading_icon(pdf, section_type, x, y, size=4.0):
    """Draw a clean, crisp vector graphic icon based on the semantic section type."""
    if section_type == "objectives":
        # Target Concentric Circles
        pdf.set_draw_color(*PDFTheme.COLOR_PRIMARY)
        pdf.set_line_width(0.4)
        pdf.circle(x + size/2, y + size/2, size/2)
        pdf.circle(x + size/2, y + size/2, size/4)
    elif section_type == "recommendations":
        # Next Steps Horizontal Arrow
        pdf.set_draw_color(*PDFTheme.COLOR_PRIMARY)
        pdf.set_line_width(0.6)
        pdf.line(x, y + size/2, x + size - 0.5, y + size/2)
        pdf.line(x + size - 2.0, y + size/2 - 1.5, x + size - 0.5, y + size/2)
        pdf.line(x + size - 2.0, y + size/2 + 1.5, x + size - 0.5, y + size/2)
    elif section_type == "risks":
        # Warning/Alert Triangle
        pdf.set_draw_color(*PDFTheme.COLOR_PRIMARY)
        pdf.set_line_width(0.4)
        p1 = (x + size/2, y)
        p2 = (x, y + size)
        p3 = (x + size, y + size)
        pdf.line(p1[0], p1[1], p2[0], p2[1])
        pdf.line(p2[0], p2[1], p3[0], p3[1])
        pdf.line(p3[0], p3[1], p1[0], p1[1])
        pdf.set_line_width(0.5)
        pdf.line(x + size/2, y + 1.2, x + size/2, y + size - 1.5)
        pdf.circle(x + size/2, y + size - 0.8, 0.2)
    elif section_type == "timeline":
        # Calendar Grid
        pdf.set_draw_color(*PDFTheme.COLOR_PRIMARY)
        pdf.set_line_width(0.4)
        pdf.rect(x, y + 0.8, size, size - 0.8, "D")
        pdf.line(x, y + 1.8, x + size, y + 1.8)
        pdf.line(x + size/3, y, x + size/3, y + 0.8)
        pdf.line(x + 2*size/3, y, x + 2*size/3, y + 0.8)
    elif section_type == "metrics":
        # Bar Chart Graph
        pdf.set_draw_color(*PDFTheme.COLOR_PRIMARY)
        pdf.set_line_width(0.4)
        pdf.line(x, y, x, y + size)
        pdf.line(x, y + size, x + size, y + size)
        pdf.set_line_width(0.6)
        pdf.line(x + 1.2, y + size, x + 1.2, y + size - 2.0)
        pdf.line(x + 2.4, y + size, x + 2.4, y + size - 3.2)
        pdf.line(x + 3.6, y + size, x + 3.6, y + size - 1.2)
    elif section_type == "deliverables":
        # Checklist Square with Tick
        pdf.set_draw_color(*PDFTheme.COLOR_PRIMARY)
        pdf.set_line_width(0.4)
        pdf.rect(x, y + 0.2, size - 0.4, size - 0.4, "D")
        pdf.set_line_width(0.5)
        pdf.line(x + 0.8, y + size/2, x + size/2.5, y + size - 1.2)
        pdf.line(x + size/2.5, y + size - 1.2, x + size - 1.0, y + 1.0)
    elif section_type == "assumptions":
        # Overlapping Rectangles (Layers)
        pdf.set_draw_color(*PDFTheme.COLOR_PRIMARY)
        pdf.set_line_width(0.4)
        pdf.rect(x, y + 1.2, size - 1.2, size - 1.2, "D")
        pdf.rect(x + 1.2, y, size - 1.2, size - 1.2, "D")


# ---------------- ACCORDION PHASE CARDS ----------------
def draw_phase_card(pdf, phase_title, lines):
    """Draw a clean, consulting-grade Phase banner card for timelines."""
    # Measure line wraps to calculate total card body height
    pdf.set_font(PDFTheme.FONT_FAMILY, "", 9)
    w_content = pdf.epw - 12  # Padding margin space
    
    total_lines = 0
    for line in lines:
        stripped = line.strip().replace('**', '').replace('*', '')
        # Check if line contains a sub-label
        parts = re.split(r'^(\*?\*?Objective:\*?\*?|\*?\*?Deliverables:\*?\*?|\*?\*?Timeline:\*?\*?|\*?\*?Budget:\*?\*?)', stripped, flags=re.IGNORECASE)
        if len(parts) >= 3:
            stripped = f"{parts[1].strip()} {parts[2].strip()}"
        l_wrap = len(pdf.multi_cell(w_content, 4.5, _to_latin1(stripped), dry_run=True, output="LINES"))
        total_lines += l_wrap
        
    body_h = (total_lines * 4.5) + (len(lines) - 1) * 2 + 8
    total_h = 8 + body_h
    
    # Check page break boundary
    pdf.check_page_break(total_h + 6)
    
    x = pdf.get_x()
    y = pdf.get_y()
    w = pdf.epw
    
    # Draw dark forest green header banner
    pdf.set_fill_color(*PDFTheme.COLOR_PRIMARY)
    pdf.rect(x, y, w, 8, "F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_font(PDFTheme.FONT_FAMILY, "B", 9.5)
    pdf.set_xy(x + 4, y + 2)
    pdf.cell(w - 8, 4, _to_latin1(phase_title.upper()))
    
    # Draw soft gray-green body fill (#F4F6F4)
    pdf.set_fill_color(*PDFTheme.COLOR_TABLE_CELL_FILL_EVEN)
    pdf.rect(x, y + 8, w, body_h, "F")
    
    # Draw thin divider border
    pdf.set_draw_color(*PDFTheme.COLOR_DIVIDER)
    pdf.set_line_width(0.3)
    pdf.rect(x, y, w, total_h, "D")
    
    current_y = y + 12
    for line in lines:
        pdf.set_xy(x + 4, current_y)
        stripped = line.strip()
        
        # Highlight sub-labels
        parts = re.split(r'^(\*?\*?Objective:\*?\*?|\*?\*?Deliverables:\*?\*?|\*?\*?Timeline:\*?\*?|\*?\*?Budget:\*?\*?)', stripped, flags=re.IGNORECASE)
        if len(parts) >= 3:
            label = parts[1].replace('**', '').replace('*', '').strip()
            val = parts[2].strip()
            
            pdf.set_font(PDFTheme.FONT_FAMILY, "B", 9)
            pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
            pdf.write(4.5, _to_latin1(f"{label} "))
            
            pdf.set_font(PDFTheme.FONT_FAMILY, "", 9)
            pdf.set_text_color(*PDFTheme.COLOR_TEXT)
            pdf.multi_cell(w - 20, 4.5, _to_latin1(val), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.set_font(PDFTheme.FONT_FAMILY, "", 9)
            pdf.set_text_color(*PDFTheme.COLOR_TEXT)
            pdf.multi_cell(w - 8, 4.5, _to_latin1(stripped.replace('**', '').replace('*', '')), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
        current_y = pdf.get_y() + 2
        
    pdf.set_y(y + total_h + 4)


# ---------------- SHADED BRIEFING PANEL (EXECUTIVE SUMMARY) ----------------
def render_executive_summary_panel(pdf, blocks):
    """Render the collected executive summary content in a single beautifully shaded panel with a left green accent spine."""
    lines = []
    for b in blocks:
        if b["type"] == "paragraph":
            lines.extend([l.strip() for l in b["content"].split("\n") if l.strip()])
        elif b["type"] == "list":
            items = parse_list_items(b["content"])
            for item in items:
                prefix = f"{item['marker']}. " if item["is_number"] else "\x95 "
                lines.append(f"{prefix}{item['text']}")
                
    if not lines:
        return
        
    pdf.set_font(PDFTheme.FONT_FAMILY, "I", 10.5)
    w_content = pdf.epw - 12
    
    total_lines = 0
    for line in lines:
        l_wrap = len(pdf.multi_cell(w_content, 5.5, _to_latin1(line), dry_run=True, output="LINES"))
        total_lines += l_wrap
        
    box_h = (total_lines * 5.5) + (len(lines) - 1) * 2 + 8
    
    # Check page overflow
    pdf.check_page_break(box_h + 6)
    
    x = pdf.get_x()
    y = pdf.get_y()
    w = pdf.epw
    
    # Draw soft light sage background (#EDF3EF)
    pdf.set_fill_color(237, 243, 239)
    pdf.rect(x, y, w, box_h, "F")
    
    # Draw left Forest Green accent strip (2mm)
    pdf.set_fill_color(*PDFTheme.COLOR_PRIMARY)
    pdf.rect(x, y, 2.0, box_h, "F")
    
    # Draw thin divider line
    pdf.set_draw_color(*PDFTheme.COLOR_DIVIDER)
    pdf.set_line_width(0.3)
    pdf.rect(x, y, w, box_h, "D")
    
    current_y = y + 4
    for line in lines:
        pdf.set_xy(x + 6, current_y)
        
        if line.startswith("\x95 "):
            pdf.set_font(PDFTheme.FONT_FAMILY, "B", 10.5)
            pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
            pdf.write(5.5, _to_latin1("\x95  "))
            pdf.set_font(PDFTheme.FONT_FAMILY, "I", 10.5)
            pdf.set_text_color(*PDFTheme.COLOR_TEXT)
            pdf.multi_cell(w_content - 4, 5.5, _to_latin1(line[2:]), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif re.match(r'^\d+\.\s+', line):
            m = re.match(r'^(\d+\.)\s+(.+)', line)
            prefix = m.group(1)
            body = m.group(2)
            pdf.set_font(PDFTheme.FONT_FAMILY, "B", 10.5)
            pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
            pdf.write(5.5, _to_latin1(f"{prefix}  "))
            pdf.set_font(PDFTheme.FONT_FAMILY, "I", 10.5)
            pdf.set_text_color(*PDFTheme.COLOR_TEXT)
            pdf.multi_cell(w_content - 6, 5.5, _to_latin1(body), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            pdf.set_font(PDFTheme.FONT_FAMILY, "I", 10.5)
            pdf.set_text_color(*PDFTheme.COLOR_TEXT)
            pdf.multi_cell(w_content, 5.5, _to_latin1(line), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
        current_y = pdf.get_y() + 2
        
    pdf.set_y(y + box_h + 6)


# ---------------- SHADED CALLOUT PANEL (OBJECTIVES / CALLOUTS) ----------------
def render_callout_box(pdf, text, border_color, fill_color, text_style=""):
    """Render a clean shaded box with a left accent border for general callouts."""
    pdf.set_font(PDFTheme.FONT_FAMILY, text_style, 10)
    w_content = pdf.epw - 12
    l_wrap = len(pdf.multi_cell(w_content, 5.5, _to_latin1(text), dry_run=True, output="LINES"))
    box_h = (l_wrap * 5.5) + 8
    
    pdf.check_page_break(box_h + 6)
    
    x = pdf.get_x()
    y = pdf.get_y()
    w = pdf.epw
    
    # Background fill
    pdf.set_fill_color(*fill_color)
    pdf.rect(x, y, w, box_h, "F")
    
    # Left border
    pdf.set_fill_color(*border_color)
    pdf.rect(x, y, 1.5, box_h, "F")
    
    # Thin outer border
    pdf.set_draw_color(*PDFTheme.COLOR_DIVIDER)
    pdf.set_line_width(0.3)
    pdf.rect(x, y, w, box_h, "D")
    
    # Text
    pdf.set_xy(x + 5, y + 4)
    pdf.set_font(PDFTheme.FONT_FAMILY, text_style, 10)
    pdf.set_text_color(*PDFTheme.COLOR_TEXT)
    pdf.multi_cell(w_content, 5.5, _to_latin1(text), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    
    pdf.set_y(y + box_h + 4)


# ---------------- RISK MATRIX RENDERER ----------------
def render_risks_table(pdf, block_content):
    """Convert risk bullets into a professional two-column mitigating table."""
    items = parse_list_items(block_content)
    table_rows = [["Risk Event / Vulnerability", "Mitigation Strategy"]]
    
    for item in items:
        text = item["text"]
        parts = re.split(r'(?i)(?:Mitigation|Response):\s*', text)
        if len(parts) >= 2:
            risk_desc = parts[0].strip().rstrip('.')
            mitigation_desc = parts[1].strip()
            table_rows.append([risk_desc, mitigation_desc])
        else:
            table_rows.append([text, "Implement standard continuous verification check."])
            
    # Set up headings Face
    from fpdf.fonts import FontFace
    heading_face = FontFace(
        fill_color=PDFTheme.COLOR_TABLE_HEADER_FILL,
        color=PDFTheme.COLOR_TABLE_HEADER_TEXT,
        emphasis="B"
    )
    pdf.set_font(PDFTheme.FONT_FAMILY, "", 8.5)
    col_widths = [pdf.epw * 0.45, pdf.epw * 0.55]
    
    pdf.check_page_break(len(table_rows) * 11 + 10)
    
    pdf.set_draw_color(*PDFTheme.COLOR_TABLE_BORDER)
    pdf.set_line_width(0.2)
    pdf.set_fill_color(*PDFTheme.COLOR_TABLE_CELL_FILL_EVEN)
    with pdf.table(
        rows=table_rows,
        align="CENTER",
        v_align="MIDDLE",
        padding=2.5,
        markdown=False,
        col_widths=col_widths,
        cell_fill_color=PDFTheme.COLOR_TABLE_CELL_FILL_EVEN,
        cell_fill_mode="EVEN_ROWS",
        headings_style=heading_face,
        min_row_height=8.0
    ) as table:
        pass
        
    pdf.set_font(PDFTheme.FONT_FAMILY, "", PDFTheme.SIZE_BODY)
    pdf.ln(4)


# ---------------- CHECKBOX BULLETS DRAWING ----------------
def draw_checkbox(pdf, x, y, checked=False, size=3.2):
    """Draw a vector checklist square box."""
    pdf.set_draw_color(*PDFTheme.COLOR_PRIMARY)
    pdf.set_line_width(0.4)
    pdf.rect(x, y, size, size, "D")
    if checked:
        pdf.set_fill_color(*PDFTheme.COLOR_PRIMARY)
        pdf.rect(x + 0.6, y + 0.6, size - 1.2, size - 1.2, "F")


# ---------------- SEMANTIC PARSING UTILITIES ----------------
def parse_markdown_to_blocks(content: str):
    """Group lines into logical markdown block objects."""
    content = content.replace("\r\n", "\n")
    raw_lines = content.split("\n")
    
    blocks = []
    current_block_type = None
    current_block_lines = []
    
    def flush_block():
        nonlocal current_block_type, current_block_lines
        if not current_block_lines:
            return
        block_text = "\n".join(current_block_lines)
        if current_block_type == "table":
            blocks.append({"type": "table", "content": block_text})
        elif current_block_type == "list":
            blocks.append({"type": "list", "content": block_text})
        elif current_block_type == "blockquote":
            blocks.append({"type": "blockquote", "content": block_text})
        else:
            stripped = block_text.strip()
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip('#'))
                blocks.append({"type": "heading", "content": stripped, "level": level})
            else:
                blocks.append({"type": "paragraph", "content": block_text})
        current_block_lines = []
        current_block_type = None

    for line in raw_lines:
        line_stripped = line.strip()
        if not line_stripped:
            flush_block()
            continue
            
        if line_stripped.startswith("|") and line_stripped.endswith("|"):
            if current_block_type != "table":
                flush_block()
                current_block_type = "table"
            current_block_lines.append(line_stripped)
        elif line_stripped.startswith(">"):
            if current_block_type != "blockquote":
                flush_block()
                current_block_type = "blockquote"
            current_block_lines.append(line_stripped)
        elif (re.match(r'^([-*+])\s+', line_stripped) or 
              re.match(r'^\d+\.\s+', line_stripped) or
              re.match(r'^\s+([-*+])\s+', line) or
              re.match(r'^\s+\d+\.\s+', line)):
            if current_block_type != "list":
                flush_block()
                current_block_type = "list"
            current_block_lines.append(line)
        elif line_stripped.startswith("#"):
            flush_block()
            current_block_lines.append(line_stripped)
            flush_block()
        else:
            if current_block_type in ["table", "list", "blockquote", "heading"]:
                flush_block()
            current_block_type = "paragraph"
            current_block_lines.append(line)
            
    flush_block()
    return blocks


def parse_list_items(block_text: str):
    """Convert a raw list block into typed structured list items."""
    items = []
    lines = block_text.split("\n")
    for line in lines:
        if not line.strip():
            continue
        indent_spaces = len(line) - len(line.lstrip())
        level = indent_spaces // 2
        line_clean = line.strip()
        
        num_match = re.match(r'^(\d+)[.)]\s+(.+)', line_clean)
        bullet_match = re.match(r'^([-*+])\s+(.+)', line_clean)
        
        if num_match:
            marker, text = num_match.groups()
            is_number = True
        elif bullet_match:
            marker, text = bullet_match.groups()
            is_number = False
        else:
            marker, text = "-", line_clean
            is_number = False
            
        checked = False
        is_checkbox = False
        if text.startswith("[ ]"):
            is_checkbox = True
            checked = False
            text = text[3:].strip()
        elif text.startswith("[x]") or text.startswith("[X]"):
            is_checkbox = True
            checked = True
            text = text[3:].strip()
            
        items.append({
            "text": text,
            "marker": marker,
            "is_number": is_number,
            "is_checkbox": is_checkbox,
            "checked": checked,
            "level": level
        })
    return items


def get_semantic_sections(content: str):
    """Classify blocks into semantic section groupings."""
    blocks = parse_markdown_to_blocks(content)
    current_section_type = "standard"
    sections = []
    
    for block in blocks:
        b_type = block["type"]
        b_content = block["content"]
        
        if b_type == "heading":
            h_text = b_content.lstrip("#").strip().replace("**", "").replace("*", "")
            h_lower = h_text.lower()
            
            if "executive summary" in h_lower or "briefing" in h_lower:
                current_section_type = "executive_summary"
            elif "objective" in h_lower or "target" in h_lower or "goal" in h_lower:
                current_section_type = "objectives"
            elif "recommendation" in h_lower or "strategy" in h_lower:
                current_section_type = "recommendations"
            elif "risk" in h_lower or "threat" in h_lower or "mitigation" in h_lower:
                current_section_type = "risks"
            elif "assumption" in h_lower or "premise" in h_lower:
                current_section_type = "assumptions"
            elif "deliverable" in h_lower or "scope" in h_lower:
                current_section_type = "deliverables"
            elif "timeline" in h_lower or "roadmap" in h_lower or "schedule" in h_lower or "phases" in h_lower:
                current_section_type = "timeline"
            elif "metric" in h_lower or "kpi" in h_lower or "roi" in h_lower or "success" in h_lower:
                current_section_type = "metrics"
            else:
                if block.get("level", 2) < 3:
                    current_section_type = "standard"
                
            sections.append({
                "type": "heading",
                "content": h_text,
                "level": len(b_content) - len(b_content.lstrip("#")),
                "section_type": current_section_type
            })
        else:
            sections.append({
                "type": b_type,
                "content": b_content,
                "section_type": current_section_type
            })
            
    return sections


# ---------------- MAIN PDF EXPORT PIPELINE ----------------
def _generate_pdf(filepath: Path, title: str, document_type: str, assumptions: list[str], content: str, goal: str, filename: str) -> str:
    import time
    from app.config import DEBUG_TIMING
    if DEBUG_TIMING:
        logger.info("[INSTRUMENTATION] PDF generation start: timestamp=%f", time.time())
        
    gen_date = datetime.now().strftime("%B %d, %Y")
    
    # Clean H1 title if present
    extracted_title = ""
    for line in content.split('\n'):
        line_strip = line.strip()
        h1_match = re.match(r'^#\s+(.+)', line_strip)
        h1_alt_match = re.match(r'^##\s+(.+)', line_strip)
        if h1_match or h1_alt_match:
            match_obj = h1_match if h1_match else h1_alt_match
            extracted_title = match_obj.group(1).strip().replace('**', '').replace('*', '')
            break
            
    if extracted_title:
        title = extracted_title
    else:
        clean_type = document_type.replace('_', ' ').title()
        if len(title) > 60 or "." in title or "," in title:
            lower_title = title.lower()
            if "onboarding" in lower_title:
                title = f"Customer Onboarding {clean_type}"
            elif "activation" in lower_title:
                title = f"Customer Activation {clean_type}"
            elif "experience" in lower_title:
                title = f"Customer Experience {clean_type}"
            elif "retention" in lower_title:
                title = f"Customer Retention {clean_type}"
            else:
                words = [w for w in title.split() if w.lower() not in ["a", "an", "the", "we", "need", "to", "create", "generate", "write", "develop", "provide", "design", "draft"]]
                short_title = " ".join(words[:4]).strip(" ,.-/\\")
                title = f"{short_title.title()} {clean_type}"
                
    pdf = PDF(doc_title=title, doc_type=document_type, gen_date=gen_date)
    pdf.set_margins(PDFTheme.MARGIN_LEFT, PDFTheme.MARGIN_TOP, PDFTheme.MARGIN_RIGHT)
    pdf.set_auto_page_break(auto=True, margin=22.0)
    
    # ---------------- 1. MCKINSEY-STYLE COVER PAGE ----------------
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)
    
    # Solid vertical accent bar (left spine)
    pdf.set_fill_color(*PDFTheme.COLOR_PRIMARY)
    pdf.rect(0, 0, 16, pdf.h, "F")
    
    # Clean background fill for remaining page (#F4F6F4)
    pdf.set_fill_color(244, 246, 244)
    pdf.rect(16, 0, pdf.w - 16, pdf.h, "F")
    
    # Cover page margins
    pdf.set_left_margin(32)
    pdf.set_right_margin(24)
    pdf.set_x(32)
    
    # Logo Monogram & Typographic Wordmark
    pdf.set_y(40)
    pdf.set_fill_color(*PDFTheme.COLOR_PRIMARY)
    pdf.rect(32, 40, 11, 11, "F")
    pdf.set_font("helvetica", "B", 6.5)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(32, 43.5)
    pdf.cell(11, 4, "AD", align="C")
    
    pdf.set_font("helvetica", "B", 10.5)
    pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
    pdf.set_xy(46, 43.5)
    pdf.cell(100, 4, "AgentDoc AI Workspace", align="L")
    
    # Document type uppercase
    pdf.set_xy(32, 90)
    pdf.set_font(PDFTheme.FONT_FAMILY, "B", 10.5)
    pdf.set_text_color(*PDFTheme.COLOR_SECONDARY)
    pdf.cell(0, 5, _to_latin1(document_type.upper().replace('_', ' ')), align="L")
    
    # Title Serif Bold (Times New Roman)
    pdf.set_xy(32, 98)
    pdf.set_font("times", "B", 26)
    pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
    pdf.multi_cell(0, 11, _to_latin1(title), align="L")
    
    # Thick Forest Green divider rule
    divider_y = pdf.get_y() + 6
    pdf.set_draw_color(*PDFTheme.COLOR_PRIMARY)
    pdf.set_line_width(1.5)
    pdf.line(32, divider_y, 85, divider_y)
    
    # Client Placeholder subtitle
    pdf.set_xy(32, divider_y + 8)
    pdf.set_font(PDFTheme.FONT_FAMILY, "I", 11)
    pdf.set_text_color(*PDFTheme.COLOR_TEXT)
    pdf.cell(0, 5, _to_latin1("Prepared for: Strategic Executive Stakeholders"), align="L")
    
    # Bottom metadata columns
    pdf.set_y(232)
    pdf.set_x(32)
    col_w = (pdf.w - 56) / 3
    
    pdf.set_font(PDFTheme.FONT_FAMILY, "B", 7.5)
    pdf.set_text_color(*PDFTheme.COLOR_SECONDARY)
    pdf.cell(col_w, 4, "CLASSIFICATION", new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.cell(col_w, 4, "PREPARED BY", new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.cell(col_w, 4, "DATE OF ISSUE", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    
    pdf.ln(1.5)
    pdf.set_x(32)
    pdf.set_font(PDFTheme.FONT_FAMILY, "", 9)
    pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
    pdf.cell(col_w, 4, "Confidential / Executive", new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.cell(col_w, 4, "AgentDoc Services", new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.cell(col_w, 4, gen_date, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    
    # Restore standard margins & auto breaks
    pdf.set_left_margin(PDFTheme.MARGIN_LEFT)
    pdf.set_right_margin(PDFTheme.MARGIN_RIGHT)
    pdf.set_auto_page_break(auto=True, margin=22.0)
    
    # ---------------- 2. CONTENT PAGES (SECTIONS RENDERING) ----------------
    pdf.add_page()
    
    # Semantic block classification
    blocks = get_semantic_sections(content)
    
    # Step A: Collect and pre-render Executive Summary blocks at the top of content page
    summary_blocks = [b for b in blocks if b["section_type"] == "executive_summary" and b["type"] != "heading"]
    other_blocks = [b for b in blocks if b["section_type"] != "executive_summary" or b["type"] == "heading"]
    
    if summary_blocks:
        pdf.set_font(PDFTheme.FONT_FAMILY, "B", PDFTheme.SIZE_H1)
        pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
        pdf.multi_cell(0, PDFTheme.LEAD_H1, _to_latin1("Executive Summary"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)
        render_executive_summary_panel(pdf, summary_blocks)
        
    # Step B: Render other blocks semantically
    idx = 0
    while idx < len(other_blocks):
        b = other_blocks[idx]
        b_type = b["type"]
        b_content = b["content"]
        sec_type = b["section_type"]
        
        # Skip empty blocks
        if not b_content.strip():
            idx += 1
            continue
            
        if b_type == "heading":
            h_text = b_content
            h_level = b["level"]
            
            # Skip double printing title heading at top of content
            if idx < 3 and (h_text.lower() in title.lower() or title.lower() in h_text.lower() or "executive summary" in h_text.lower()):
                idx += 1
                continue
                
            if h_level == 1:
                pdf.check_page_break(25)
                # Section horizontal rule divider
                if pdf.page_no() > 2 or (pdf.page_no() == 2 and pdf.get_y() > pdf.t_margin + 15):
                    pdf.ln(3)
                    pdf.set_draw_color(*PDFTheme.COLOR_DIVIDER)
                    pdf.set_line_width(0.3)
                    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
                    pdf.ln(5)
                
                # Draw Vector Icon next to section heading
                icon_y = pdf.get_y() + 1.2
                draw_heading_icon(pdf, sec_type, pdf.l_margin, icon_y, size=4.2)
                
                pdf.set_left_margin(PDFTheme.MARGIN_LEFT + 7.5)
                pdf.set_x(PDFTheme.MARGIN_LEFT + 7.5)
                pdf.set_font(PDFTheme.FONT_FAMILY, "B", PDFTheme.SIZE_H1)
                pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
                pdf.multi_cell(0, PDFTheme.LEAD_H1, _to_latin1(h_text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.set_left_margin(PDFTheme.MARGIN_LEFT)
                pdf.ln(4)
                
            elif h_level in [2, 3] and sec_type == "timeline" and h_text.lower().startswith("phase"):
                # Check for Phase card timeline groupings
                # Look ahead: if next block is a List block, combine them into a Phase Card!
                phase_lines = []
                if idx + 1 < len(other_blocks) and other_blocks[idx + 1]["type"] == "list":
                    list_b = other_blocks[idx + 1]
                    phase_lines = [l.strip() for l in list_b["content"].split("\n") if l.strip()]
                    idx += 1 # Consume list block
                else:
                    phase_lines = ["Objective: Execute defined timeline roadmap steps."]
                
                draw_phase_card(pdf, h_text, phase_lines)
                idx += 1
                continue
                
            elif h_level == 2:
                pdf.check_page_break(18)
                pdf.set_font(PDFTheme.FONT_FAMILY, "B", PDFTheme.SIZE_H2)
                pdf.set_text_color(*PDFTheme.COLOR_SECONDARY)
                pdf.multi_cell(0, PDFTheme.LEAD_H2, _to_latin1(h_text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(3)
                
            else:
                pdf.check_page_break(14)
                pdf.set_font(PDFTheme.FONT_FAMILY, "B", PDFTheme.SIZE_H3)
                pdf.set_text_color(*PDFTheme.COLOR_SECONDARY)
                pdf.multi_cell(0, PDFTheme.LEAD_H3, _to_latin1(h_text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2.5)
                
        elif b_type == "list":
            # Semantic risks matrix redirection
            if sec_type == "risks":
                render_risks_table(pdf, b_content)
                idx += 1
                continue
                
            items = parse_list_items(b_content)
            for item in items:
                pdf.check_page_break(11)
                
                # Dynamic list item indentation
                list_indent = PDFTheme.MARGIN_LEFT + (item["level"] * 6) + 6
                marker_pos = PDFTheme.MARGIN_LEFT + (item["level"] * 6)
                
                pdf.set_left_margin(list_indent)
                pdf.set_x(marker_pos)
                
                if item["is_checkbox"] or sec_type in ["assumptions", "deliverables"]:
                    # Draw visual square checklist box
                    draw_checkbox(pdf, marker_pos, pdf.get_y() + 1.2, checked=item["checked"] or sec_type == "deliverables", size=3.2)
                    pdf.set_x(list_indent)
                else:
                    pdf.set_font(PDFTheme.FONT_FAMILY, "B", PDFTheme.SIZE_BODY)
                    pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
                    if item["is_number"]:
                        pdf.write(PDFTheme.LEAD_BODY, _to_latin1(f"{item['marker']}. "))
                    else:
                        pdf.write(PDFTheme.LEAD_BODY, _to_latin1("\x95  "))
                        
                pdf.set_font(PDFTheme.FONT_FAMILY, "", PDFTheme.SIZE_BODY)
                pdf.set_text_color(*PDFTheme.COLOR_TEXT)
                pdf.multi_cell(0, PDFTheme.LEAD_BODY, _to_latin1(item["text"]), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.set_left_margin(PDFTheme.MARGIN_LEFT)
                pdf.ln(1)
                
            pdf.ln(2)
            
        elif b_type == "table":
            # Standard Markdown Table parser
            table_rows = []
            for row_line in b_content.split("\n"):
                row_line = row_line.strip()
                if row_line.startswith("|") and row_line.endswith("|"):
                    if not re.match(r'^\|[\s\-:|]+\|$', row_line):
                        cells = [c.strip() for c in row_line.strip("|").split("|")]
                        if cells and cells[0] == "":
                            cells = cells[1:]
                        if cells and cells[-1] == "":
                            cells = cells[:-1]
                        table_rows.append(cells)
            
            if table_rows:
                try:
                    from fpdf.fonts import FontFace
                    heading_face = FontFace(
                        fill_color=PDFTheme.COLOR_TABLE_HEADER_FILL,
                        color=PDFTheme.COLOR_TABLE_HEADER_TEXT,
                        emphasis="B"
                    )
                    
                    cleaned_rows = []
                    for row in table_rows:
                        cleaned_row = []
                        for cell in row:
                            cell_cleaned = _clean_all_markdown_for_tables(cell)
                            cell_cleaned = _to_latin1(cell_cleaned)
                            cleaned_row.append(cell_cleaned)
                        cleaned_rows.append(cleaned_row)
                        
                    pdf.set_font(PDFTheme.FONT_FAMILY, "", 8.5)
                    pdf.set_text_color(*PDFTheme.COLOR_TEXT)
                    
                    num_cols = len(cleaned_rows[0])
                    cell_padding = 2.5
                    
                    col_max_widths = [0.0] * num_cols
                    col_max_word_widths = [0.0] * num_cols
                    for row in cleaned_rows:
                        for col_idx, cell in enumerate(row):
                            if col_idx < num_cols:
                                str_w = pdf.get_string_width(cell)
                                col_max_widths[col_idx] = max(col_max_widths[col_idx], str_w)
                                words = cell.split()
                                word_w = max(pdf.get_string_width(w) for w in words) if words else 0.0
                                col_max_word_widths[col_idx] = max(col_max_word_widths[col_idx], word_w)
                                
                    total_text_w = sum(col_max_widths)
                    usable_width = pdf.epw
                    
                    if total_text_w > 0:
                        col_widths = [(w / total_text_w) * usable_width for w in col_max_widths]
                        for col_i in range(num_cols):
                            word_limit = col_max_word_widths[col_i] + (2 * cell_padding) + 1.0
                            hard_min = max(25.0, usable_width * 0.12)
                            col_min = max(word_limit, hard_min)
                            if col_widths[col_i] < col_min:
                                col_widths[col_i] = col_min
                        total_w = sum(col_widths)
                        if total_w > usable_width:
                            col_widths = [(w / total_w) * usable_width for w in col_widths]
                        else:
                            remaining = usable_width - total_w
                            if remaining > 0:
                                col_widths = [w + (col_max_widths[col_i] / total_text_w) * remaining for col_i, w in enumerate(col_widths)]
                    else:
                        col_widths = None
                        
                    estimated_height = len(cleaned_rows) * 11
                    pdf.check_page_break(estimated_height if estimated_height < 60 else 30)
                    
                    pdf.set_draw_color(*PDFTheme.COLOR_TABLE_BORDER)
                    pdf.set_line_width(0.2)
                    pdf.set_fill_color(*PDFTheme.COLOR_TABLE_CELL_FILL_EVEN)
                    with pdf.table(
                        rows=cleaned_rows,
                        align="CENTER",
                        v_align="MIDDLE",
                        padding=cell_padding,
                        markdown=False,
                        col_widths=col_widths,
                        cell_fill_color=PDFTheme.COLOR_TABLE_CELL_FILL_EVEN,
                        cell_fill_mode="EVEN_ROWS",
                        headings_style=heading_face,
                        min_row_height=8.0
                    ) as table:
                        pass
                    
                    pdf.set_font(PDFTheme.FONT_FAMILY, "", PDFTheme.SIZE_BODY)
                    pdf.ln(4)
                except Exception as ex:
                    logger.error("Error rendering table to PDF: %s", ex)
                    
        elif b_type == "blockquote":
            # Clean blockquote marker '>'
            text = b_content.lstrip(">").strip()
            render_callout_box(pdf, text, PDFTheme.COLOR_SECONDARY, (244, 246, 244), text_style="I")
            
        else:
            # Semantic Objectives Redirect
            if sec_type == "objectives" and len(b_content) > 30:
                render_callout_box(pdf, b_content, PDFTheme.COLOR_PRIMARY, (244, 246, 244))
                idx += 1
                continue
                
            # Paragraph Block Rendering
            lines = [l.strip() for l in b_content.split("\n") if l.strip()]
            for line in lines:
                pdf.check_page_break(12)
                
                # Check for inline bullet formatting
                if line.startswith("- ") or line.startswith("* "):
                    pdf.set_left_margin(PDFTheme.MARGIN_LEFT + 6)
                    pdf.set_x(PDFTheme.MARGIN_LEFT)
                    pdf.set_font(PDFTheme.FONT_FAMILY, "B", PDFTheme.SIZE_BODY)
                    pdf.set_text_color(*PDFTheme.COLOR_PRIMARY)
                    pdf.write(PDFTheme.LEAD_BODY, _to_latin1("\x95  "))
                    pdf.set_font(PDFTheme.FONT_FAMILY, "", PDFTheme.SIZE_BODY)
                    pdf.set_text_color(*PDFTheme.COLOR_TEXT)
                    pdf.multi_cell(0, PDFTheme.LEAD_BODY, _to_latin1(line[2:]), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.set_left_margin(PDFTheme.MARGIN_LEFT)
                    pdf.ln(1)
                else:
                    pdf.set_font(PDFTheme.FONT_FAMILY, "", PDFTheme.SIZE_BODY)
                    pdf.set_text_color(*PDFTheme.COLOR_TEXT)
                    pdf.multi_cell(0, PDFTheme.LEAD_BODY, _to_latin1(line), markdown=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.ln(2.5)
            pdf.ln(1.5)
            
        idx += 1
        
    if DEBUG_TIMING:
        logger.info("[INSTRUMENTATION] PDF save/write start: timestamp=%f", time.time())
        
    pdf.output(str(filepath))
    
    return filename


def _clean_markdown_preserving_formatting(text: str) -> str:
    """Remove code blocks and links, but preserve bold and italic markers."""
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text.strip()


def _clean_all_markdown_for_tables(text: str) -> str:
    """Remove code blocks, links, bold, and italic markers completely for tables."""
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = text.replace('**', '').replace('*', '')
    return text.strip()


def _break_long_words(text: str, max_length: int = 50) -> str:
    """Break long sequences of non-whitespace characters to prevent PDF rendering errors."""
    return re.sub(r'(\S{' + str(max_length) + r'})', r'\1 ', text)


def _add_content_sections(doc: Document, content: str) -> None:
    # 1. Parse content using get_semantic_sections
    blocks = get_semantic_sections(content)
    
    # 2. Extract Executive Summary blocks
    summary_blocks = [b for b in blocks if b["section_type"] == "executive_summary" and b["type"] != "heading"]
    other_blocks = [b for b in blocks if b["section_type"] != "executive_summary" or b["type"] == "heading"]
    
    # 3. Render Executive Summary (if present) inside the special panel
    if summary_blocks:
        p_es = doc.add_paragraph()
        p_es.paragraph_format.space_before = Pt(12)
        p_es.paragraph_format.space_after = Pt(6)
        p_es.paragraph_format.keep_with_next = True
        run = p_es.add_run("Executive Summary")
        run.bold = True
        run.font.name = DOCXTheme.FONT_BODY
        run.font.size = Pt(DOCXTheme.SIZE_H1)
        run.font.color.rgb = DOCXTheme.RGB_PRIMARY
        
        _docx_render_executive_summary_panel(doc, summary_blocks)
        
    # 4. Render other blocks semantically
    idx = 0
    while idx < len(other_blocks):
        b = other_blocks[idx]
        b_type = b["type"]
        b_content = b["content"]
        sec_type = b["section_type"]
        
        if not b_content.strip():
            idx += 1
            continue
            
        if b_type == "heading":
            h_text = b_content
            h_level = b["level"]
            
            # Skip duplicate title headers at the top
            if idx < 3 and "executive summary" in h_text.lower():
                idx += 1
                continue
                
            if h_level == 1:
                p_head = doc.add_paragraph()
                p_head.paragraph_format.space_before = Pt(18)
                p_head.paragraph_format.space_after = Pt(8)
                p_head.paragraph_format.keep_with_next = True
                
                # Apply top border as section divider line
                if idx > 2:
                    _set_paragraph_top_border(p_head, hex_color=DOCXTheme.HEX_DIVIDER, size=4)
                    
                run = p_head.add_run(h_text)
                run.bold = True
                run.font.name = DOCXTheme.FONT_BODY
                run.font.size = Pt(DOCXTheme.SIZE_H1)
                run.font.color.rgb = DOCXTheme.RGB_PRIMARY
                
            elif h_level in [2, 3] and sec_type == "timeline" and h_text.lower().startswith("phase"):
                # Phase Timeline card lookahead
                phase_lines = []
                if idx + 1 < len(other_blocks) and other_blocks[idx + 1]["type"] == "list":
                    list_b = other_blocks[idx + 1]
                    phase_lines = [l.strip() for l in list_b["content"].split("\n") if l.strip()]
                    idx += 1 # Consume list block
                else:
                    phase_lines = ["Objective: Execute defined timeline roadmap steps."]
                    
                _docx_render_phase_card(doc, h_text, phase_lines)
                idx += 1
                continue
                
            elif h_level == 2:
                p_head = doc.add_paragraph()
                p_head.paragraph_format.space_before = Pt(14)
                p_head.paragraph_format.space_after = Pt(6)
                p_head.paragraph_format.keep_with_next = True
                run = p_head.add_run(h_text)
                run.bold = True
                run.font.name = DOCXTheme.FONT_BODY
                run.font.size = Pt(DOCXTheme.SIZE_H2)
                run.font.color.rgb = DOCXTheme.RGB_SECONDARY
                
            else:
                p_head = doc.add_paragraph()
                p_head.paragraph_format.space_before = Pt(10)
                p_head.paragraph_format.space_after = Pt(4)
                p_head.paragraph_format.keep_with_next = True
                run = p_head.add_run(h_text)
                run.bold = True
                run.font.name = DOCXTheme.FONT_BODY
                run.font.size = Pt(DOCXTheme.SIZE_H3)
                run.font.color.rgb = DOCXTheme.RGB_SECONDARY
                
        elif b_type == "list":
            if sec_type == "risks":
                _docx_render_risk_matrix(doc, b_content)
                idx += 1
                continue
                
            items = parse_list_items(b_content)
            for item in items:
                style_name = 'List Bullet'
                prefix = ""
                if item["is_checkbox"] or sec_type in ["assumptions", "deliverables"]:
                    style_name = 'Normal'
                    box_char = "☑  " if item["checked"] or sec_type == "deliverables" else "☐  "
                    prefix = f"{box_char}"
                elif item["is_number"]:
                    style_name = 'List Number'
                    
                p_li = doc.add_paragraph(style=style_name)
                p_li.paragraph_format.space_before = Pt(0)
                p_li.paragraph_format.space_after = Pt(3)
                p_li.paragraph_format.line_spacing = 1.15
                
                indent_pt = 18 + (item["level"] * 18)
                p_li.paragraph_format.left_indent = Pt(indent_pt)
                
                if prefix:
                    run_pref = p_li.add_run(prefix)
                    run_pref.font.name = DOCXTheme.FONT_BODY
                    run_pref.font.size = Pt(DOCXTheme.SIZE_BODY)
                    run_pref.font.color.rgb = DOCXTheme.RGB_PRIMARY
                    run_pref.bold = True
                    
                text = _clean_markdown(item["text"])
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p_li.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p_li.add_run(part)
                    run.font.name = DOCXTheme.FONT_BODY
                    run.font.size = Pt(DOCXTheme.SIZE_BODY)
                    run.font.color.rgb = DOCXTheme.RGB_TEXT
                    
        elif b_type == "table":
            table_rows = []
            for row_line in b_content.split("\n"):
                row_line = row_line.strip()
                if row_line.startswith("|") and row_line.endswith("|"):
                    if not re.match(r'^\|[\s\-:|]+\|$', row_line):
                        cells = [c.strip() for c in row_line.strip("|").split("|")]
                        if cells and cells[0] == "":
                            cells = cells[1:]
                        if cells and cells[-1] == "":
                            cells = cells[:-1]
                        table_rows.append(cells)
            if table_rows:
                _docx_render_consulting_table(doc, table_rows)
                
        elif b_type == "blockquote":
            text = b_content.lstrip(">").strip()
            _docx_render_callout_box(doc, text, DOCXTheme.HEX_SECONDARY, DOCXTheme.HEX_BG_LIGHT, text_style="I")
            
        else:
            if sec_type == "objectives" and len(b_content) > 30:
                _docx_render_callout_box(doc, b_content, DOCXTheme.HEX_PRIMARY, DOCXTheme.HEX_BG_LIGHT)
                idx += 1
                continue
                
            lines = [l.strip() for l in b_content.split("\n") if l.strip()]
            for line in lines:
                if line.startswith("- ") or line.startswith("* "):
                    p_li = doc.add_paragraph(style='List Bullet')
                    p_li.paragraph_format.space_before = Pt(0)
                    p_li.paragraph_format.space_after = Pt(3)
                    p_li.paragraph_format.line_spacing = 1.15
                    p_li.paragraph_format.left_indent = Pt(18)
                    
                    text = _clean_markdown(line[2:])
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p_li.add_run(part[2:-2])
                            run.bold = True
                        else:
                            run = p_li.add_run(part)
                        run.font.name = DOCXTheme.FONT_BODY
                        run.font.size = Pt(DOCXTheme.SIZE_BODY)
                        run.font.color.rgb = DOCXTheme.RGB_TEXT
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.line_spacing = 1.15
                    
                    text = _clean_markdown(line)
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            run = p.add_run(part)
                        run.font.name = DOCXTheme.FONT_BODY
                        run.font.size = Pt(DOCXTheme.SIZE_BODY)
                        run.font.color.rgb = DOCXTheme.RGB_TEXT
                        
        idx += 1


def _clean_markdown(text: str) -> str:
    """Remove common markdown formatting from text, preserving bold."""
    # Remove inline code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove markdown links, keep text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text.strip()


def _editorial_cleanup_pass(text: str) -> str:
    """Lightweight editorial pass to fix formatting, spacing, and punctuation issues."""
    if not text:
        return ""
    # 1. Remove repeated periods
    text = re.sub(r'\.{2,}', '.', text)
    # 2. Remove duplicate commas
    text = re.sub(r',{2,}', ',', text)
    # 3. Correct spacing around punctuation: space after punctuation, no space before
    text = re.sub(r'\s+([.,;:!?])', r'\1', text)
    # Ensure space after punctuation (but not if followed by a number or punctuation)
    text = re.sub(r'([.,;:!?])(?=[A-Za-z])', r'\1 ', text)
    # 4. Clean up typical spelling or spacing bugs
    text = re.sub(r'\bcore\s+issueunknown\b', 'core issue: unknown', text)
    # 5. Ensure headings don't have trailing punctuation like periods
    lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("#"):
            if stripped.endswith("."):
                line = line.rstrip(".")
        lines.append(line)
    text = "\n".join(lines)
    return text
